"""Build a polished WORKFLOW.docx from the content in WORKFLOW.md.

This script is the reproducible source for ``WORKFLOW.docx`` at the repo root.
It renders two diagrams with matplotlib (an architecture diagram and a
stage-by-stage workflow flow diagram), then assembles a styled Word document
with python-docx: a title page, a table of contents, colour-coded headings,
shaded callout boxes, and well-formatted tables.

Run:  python tools/build_workflow_docx.py
Output: WORKFLOW.docx (and tools/_assets/*.png intermediate images)
"""
from __future__ import annotations

import os

import matplotlib

matplotlib.use("Agg")
import matplotlib.patches as mpatches
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# --------------------------------------------------------------------------
# Paths & palette
# --------------------------------------------------------------------------
REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ASSETS = os.path.join(REPO, "tools", "_assets")
os.makedirs(ASSETS, exist_ok=True)
ARCH_PNG = os.path.join(ASSETS, "architecture.png")
FLOW_PNG = os.path.join(ASSETS, "workflow_flow.png")
DOCX_PATH = os.path.join(REPO, "WORKFLOW.docx")

# Finance blue-green palette (hex without leading #)
INK = "0F2E3D"        # near-black teal (body-ish dark)
PRIMARY = "0F4C5C"    # deep teal  (H1 / title)
SECONDARY = "1B6B7A"  # teal       (H2)
ACCENT = "2A9D8F"     # green-teal (H3 / accents)
GOLD = "E9A23B"       # warm accent for callouts
LIGHT = "E8F1F2"      # very light bg for callout
MINT = "DCF0EC"       # light mint for info callout
SAND = "FBEFD9"       # light sand for warning callout
TABLE_HDR = "0F4C5C"  # table header fill
ROW_ALT = "EAF3F4"    # alt row fill

# matplotlib equivalents (with #)
M_PRIMARY = "#0F4C5C"
M_SECONDARY = "#1B6B7A"
M_ACCENT = "#2A9D8F"
M_GOLD = "#E9A23B"
M_LIGHT = "#E8F1F2"
M_EXT = "#5B7DB1"
M_DATA = "#7C6BA0"
M_WHITE = "#FFFFFF"
M_INK = "#0F2E3D"


# ==========================================================================
# Diagram rendering (matplotlib)
# ==========================================================================
def _box(ax, x, y, w, h, text, face, edge, tcolor="#FFFFFF", fs=10, bold=True,
         rounded=0.08):
    patch = FancyBboxPatch(
        (x, y), w, h,
        boxstyle=f"round,pad=0.02,rounding_size={rounded}",
        linewidth=1.4, edgecolor=edge, facecolor=face, zorder=2,
    )
    ax.add_patch(patch)
    ax.text(
        x + w / 2, y + h / 2, text, ha="center", va="center",
        fontsize=fs, color=tcolor, weight="bold" if bold else "normal",
        zorder=3, wrap=True,
    )
    return (x + w / 2, y + h / 2)


def _arrow(ax, p1, p2, color=M_INK, style="-|>", lw=1.6, ls="-"):
    ax.add_patch(FancyArrowPatch(
        p1, p2, arrowstyle=style, mutation_scale=14,
        color=color, linewidth=lw, linestyle=ls,
        shrinkA=6, shrinkB=6, zorder=1,
    ))


def render_architecture():
    """Component architecture: user -> CLI -> orchestrator -> agent pipeline -> outputs,
    with external data sources, LLM provider and profiles feeding in."""
    fig, ax = plt.subplots(figsize=(11.5, 7.4), dpi=170)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    # Top row: user -> cli -> orchestrator
    c_user = _box(ax, 3, 86, 18, 9, "You\n(CLI command)", M_GOLD, "#B97D1E",
                  tcolor=M_INK, fs=10)
    c_cli = _box(ax, 28, 86, 20, 9, "src/cli.py\nparses options", M_SECONDARY,
                 M_PRIMARY)
    c_orch = _box(ax, 55, 86, 24, 9, "src/graph/orchestrator.py\nruns the pipeline",
                  M_PRIMARY, M_INK)
    _arrow(ax, (21, 90.5), (28, 90.5))
    _arrow(ax, (48, 90.5), (55, 90.5))

    # Pipeline container
    ax.add_patch(FancyBboxPatch(
        (3, 20), 60, 60, boxstyle="round,pad=0.02,rounding_size=0.6",
        linewidth=1.6, edgecolor=M_PRIMARY, facecolor="#F4F9FA", zorder=0))
    ax.text(20, 77.5, "The Pipeline of Agents", ha="left", va="center",
            fontsize=12, weight="bold", color=M_PRIMARY)

    stages = [
        ("Universe Builder\nuniverse.py", M_ACCENT),
        ("Quant + Factor Engine\nquant.py", M_ACCENT),
        ("Fundamentals\nfundamentals.py", M_SECONDARY),
        ("Technical\ntechnical.py", M_SECONDARY),
        ("News + Sentiment\nnews_sentiment.py", M_SECONDARY),
        ("Macro\nmacro.py", M_SECONDARY),
        ("Bull vs Bear Debate\nresearchers.py", M_GOLD),
        ("Risk + Tax\nrisk_profile.py", M_ACCENT),
        ("Research Manager\nmanager.py", M_PRIMARY),
    ]
    bx, bw, bh = 12, 42, 4.6
    top = 71
    gap = 5.85
    centers = []
    for i, (label, color) in enumerate(stages):
        y = top - i * gap
        tcol = M_INK if color == M_GOLD else M_WHITE
        c = _box(ax, bx, y, bw, bh, label, color, M_INK, tcolor=tcol, fs=9,
                 rounded=0.05)
        centers.append((bx, y, bw, bh, c))
        if i > 0:
            px, py, pw, ph, pc = centers[i - 1]
            _arrow(ax, (bx + bw / 2, py), (bx + bw / 2, y + bh),
                   color=M_PRIMARY, lw=1.3)

    # orchestrator -> pipeline
    _arrow(ax, (62, 86), (52, 73), color=M_PRIMARY, lw=1.6)

    # Right column: report / memory / outputs
    c_rep = _box(ax, 72, 60, 25, 8, "Report Generator\nreport/generator.py",
                 M_PRIMARY, M_INK, fs=9)
    c_mem = _box(ax, 72, 48, 25, 8, "Memory Log\nmemory/store.py",
                 M_SECONDARY, M_PRIMARY, fs=9)
    c_out = _box(ax, 72, 35, 25, 9,
                 "Outputs\nreports/*.md  +  *.xlsx", M_GOLD, "#B97D1E",
                 tcolor=M_INK, fs=9)
    # manager -> report
    mgr_c = centers[-1][4]
    _arrow(ax, (54, mgr_c[1]), (72, 64), color=M_PRIMARY, lw=1.6)
    _arrow(ax, (84.5, 60), (84.5, 56), color=M_PRIMARY)
    _arrow(ax, (84.5, 48), (84.5, 44), color=M_PRIMARY)

    # External inputs (bottom)
    c_data = _box(ax, 3, 5, 27, 10,
                  "Free Data Sources\nyfinance · Stooq\nGDELT · SEC EDGAR",
                  M_DATA, "#4A3F66", fs=9)
    c_llm = _box(ax, 36.5, 5, 27, 10,
                 "LLM Provider\nOpenAI · Anthropic\nOllama · stub",
                 M_EXT, "#3C557D", fs=9)
    c_prof = _box(ax, 70, 5, 27, 10,
                  "config/profiles/*.yaml\nweights · tax · risk",
                  "#9C8FB8", "#4A3F66", tcolor=M_INK, fs=9)
    # data feeds quant + news
    _arrow(ax, (16, 15), (20, top - 1 * gap), color=M_DATA, ls="--", lw=1.2)
    _arrow(ax, (20, 15), (24, top - 4 * gap + 2), color=M_DATA, ls="--", lw=1.2)
    # llm feeds analysts/debate/manager
    _arrow(ax, (50, 15), (40, top - 6 * gap + 2), color=M_EXT, ls="--", lw=1.2)
    # profiles feed cli
    _arrow(ax, (83, 15), (40, 86), color="#9C8FB8", ls="--", lw=1.2)

    # Layer legend
    legend_items = [
        ("Perception (get data)", M_DATA),
        ("Brain: deterministic math", M_ACCENT),
        ("Brain: AI interpretation", M_SECONDARY),
        ("Action (write output)", M_PRIMARY),
    ]
    handles = [mpatches.Patch(color=c, label=l) for l, c in legend_items]
    ax.legend(handles=handles, loc="upper right", bbox_to_anchor=(1.0, 1.0),
              fontsize=8.5, frameon=True, title="Three layers",
              title_fontsize=9)

    plt.tight_layout()
    fig.savefig(ARCH_PNG, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def render_workflow_flow():
    """Stage-by-stage flow diagram (Stage 0 -> 8) with the LLM decision branch."""
    fig, ax = plt.subplots(figsize=(11.5, 8.2), dpi=170)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    def stage(x, y, w, h, num, title, color, tcol=M_WHITE, fs=9.2):
        ax.add_patch(FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.4",
            linewidth=1.4, edgecolor=M_INK, facecolor=color, zorder=2))
        ax.text(x + 3.0, y + h / 2, num, ha="center", va="center",
                fontsize=13, weight="bold", color=tcol, zorder=3)
        ax.text(x + w / 2 + 2.5, y + h / 2, title, ha="center", va="center",
                fontsize=fs, color=tcol, weight="bold", zorder=3)
        return (x + w / 2, y, x + w / 2, y + h, x + w, y + h / 2, x, y + h / 2)

    w, h = 40, 7.2
    cx = 8
    ys = [88, 77.5, 67, 56.5]
    s0 = stage(cx, ys[0], w, h, "0", "Run a command (cli.py)", M_GOLD, M_INK)
    s1 = stage(cx, ys[1], w, h, "1", "Universe Builder", M_ACCENT)
    s2 = stage(cx, ys[2], w, h, "2", "Quant + Factor Engine", M_PRIMARY)
    dec_y = 47
    # decision diamond
    dcx, dcy = cx + w / 2, dec_y + 3.5
    diamond = mpatches.FancyBboxPatch(
        (dcx - 13, dec_y), 26, 7, boxstyle="round,pad=0.02,rounding_size=3.5",
        linewidth=1.4, edgecolor="#B97D1E", facecolor=M_GOLD, zorder=2)
    ax.add_patch(diamond)
    ax.text(dcx, dcy, "should_run_llm?", ha="center", va="center",
            fontsize=9.2, weight="bold", color=M_INK, zorder=3)

    # connect 0->1->2->decision
    _arrow(ax, (s0[0], s0[1]), (s1[2], s1[3]), color=M_PRIMARY)
    _arrow(ax, (s1[0], s1[1]), (s2[2], s2[3]), color=M_PRIMARY)
    _arrow(ax, (s2[0], s2[1]), (dcx, dec_y + 7), color=M_PRIMARY)

    # LEFT branch: AI path (stages 3,4)
    aiw = 30
    aix = 2
    a3 = stage(aix, 33, aiw, 6.4, "3", "AI Analysts (x4)", M_SECONDARY, fs=8.6)
    a4 = stage(aix, 24, aiw, 6.4, "4", "Bull vs Bear Debate", M_SECONDARY, fs=8.6)
    _arrow(ax, (dcx - 13, dcy), (aix + aiw / 2, 39.4), color=M_ACCENT)
    ax.text(20, 44.5, "LLM enabled\n& shortlist", ha="center", va="center",
            fontsize=7.8, color=M_ACCENT, weight="bold")
    _arrow(ax, (a3[0], a3[1]), (a4[2], a4[3]), color=M_SECONDARY)

    # RIGHT branch: quant-only path
    qw = 30
    qx = 56
    q = stage(qx, 30, qw, 7.0, "", "Manager.run_quant_only", "#9C8FB8",
              tcol=M_INK, fs=8.4)
    _arrow(ax, (dcx + 13, dcy), (qx + qw / 2, 37), color="#B97D1E")
    ax.text(67, 44, "--no-llm\nor empty", ha="center", va="center",
            fontsize=7.8, color="#B97D1E", weight="bold")

    # Merge into stage 5,6,7,8 (centered bottom)
    s5 = stage(cx, 14.5, w, 6.6, "5", "Risk + Profile / Tax", M_ACCENT, fs=8.8)
    s6 = stage(cx + 46, 14.5, 42, 6.6, "6", "Research Manager (picks)",
               M_PRIMARY, fs=8.8)
    s7 = stage(cx, 5, w, 6.6, "7", "Report Generator (.md/.xlsx)", M_PRIMARY,
               fs=8.6)
    s8 = stage(cx + 46, 5, 42, 6.6, "8", "Memory Log", M_SECONDARY, fs=8.8)

    # AI path -> stage 5 ; quant -> stage 5
    _arrow(ax, (a4[0], a4[1]), (cx + 6, 21.1), color=M_SECONDARY)
    _arrow(ax, (qx + qw / 2, 30), (cx + w - 4, 21.1), color="#9C8FB8")
    # 5 -> 6 -> 7 -> 8 (snake order)
    _arrow(ax, (s5[4], s5[5]), (s6[6], s6[5]), color=M_PRIMARY)
    _arrow(ax, (cx + 46 + 8, 14.5), (cx + w - 4, 11.6), color=M_PRIMARY)
    _arrow(ax, (s7[4], s7[5]), (s8[6], s8[5]), color=M_PRIMARY)

    ax.text(50, 97, "End-to-End Workflow: from typed command to report on disk",
            ha="center", va="center", fontsize=12.5, weight="bold",
            color=M_PRIMARY)

    plt.tight_layout()
    fig.savefig(FLOW_PNG, bbox_inches="tight", facecolor="white")
    plt.close(fig)


# ==========================================================================
# python-docx helpers
# ==========================================================================
def _shade(cell, hex_fill):
    """Apply background shading to a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for tag, val in (("top", top), ("bottom", bottom),
                     ("start", left), ("end", right)):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _no_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        borders.append(e)
    tblPr.append(borders)


def _set_table_borders(table, color="BFD7DB", sz=4):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def _runfmt(run, size=None, bold=None, italic=None, color=None, font="Calibri"):
    run.font.name = font
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_callout(doc, title, body_lines, fill, bar_color, icon=""):
    """A shaded one-cell table acting as a callout box, with an accent bar."""
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _no_table_borders(table)
    table.columns[0].width = Inches(0.12)
    table.columns[1].width = Inches(6.3)

    bar = table.rows[0].cells[0]
    _shade(bar, bar_color)
    bar.width = Inches(0.12)

    cell = table.rows[0].cells[1]
    _shade(cell, fill)
    _set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    cell.width = Inches(6.3)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{icon}  {title}" if icon else title)
    _runfmt(r, size=11.5, bold=True, color=bar_color)

    for line in body_lines:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_after = Pt(2)
        if isinstance(line, list):
            for seg_text, seg_bold in line:
                rr = bp.add_run(seg_text)
                _runfmt(rr, size=10.5, bold=seg_bold, color=INK)
        else:
            rr = bp.add_run(line)
            _runfmt(rr, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def styled_table(doc, headers, rows, col_widths=None, header_fill=TABLE_HDR,
                 font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_table_borders(table)

    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        _shade(hdr[i], header_fill)
        _set_cell_margins(hdr[i])
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(htext)
        _runfmt(r, size=font_size + 0.5, bold=True, color="FFFFFF")

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            _set_cell_margins(cells[i])
            if ridx % 2 == 1:
                _shade(cells[i], ROW_ALT)
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            # support bold first segment via tuple
            if isinstance(val, tuple):
                r = p.add_run(val[0])
                _runfmt(r, size=font_size, bold=True, color=PRIMARY)
                if len(val) > 1 and val[1]:
                    r2 = p.add_run(val[1])
                    _runfmt(r2, size=font_size, color=INK)
            else:
                r = p.add_run(str(val))
                _runfmt(r, size=font_size, color=INK)
    if col_widths:
        for i, wdt in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(wdt)
    return table


def add_heading(doc, text, level):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(14 if level <= 1 else 10)
    h.paragraph_format.space_after = Pt(6)
    r = h.add_run(text)
    if level == 1:
        _runfmt(r, size=18, bold=True, color=PRIMARY, font="Calibri")
    elif level == 2:
        _runfmt(r, size=14, bold=True, color=SECONDARY, font="Calibri")
    else:
        _runfmt(r, size=12, bold=True, color=ACCENT, font="Calibri")
    return h


def add_body(doc, text, size=10.5, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    _runfmt(r, size=size, italic=italic, color=INK)
    return p


def add_bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, list):
            for seg_text, seg_bold in it:
                r = p.add_run(seg_text)
                _runfmt(r, size=size, bold=seg_bold, color=INK)
        else:
            r = p.add_run(it)
            _runfmt(r, size=size, color=INK)


def add_image_with_caption(doc, path, caption, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    r = cap.add_run(caption)
    _runfmt(r, size=9, italic=True, color=SECONDARY)


def add_toc_field(doc):
    """Insert a real Word TOC field (updates on open / F9)."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldBegin = OxmlElement("w:fldChar")
    fldBegin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fldSep = OxmlElement("w:fldChar")
    fldSep.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Right-click and choose \u201cUpdate Field\u201d to build the table of contents."
    fldEnd = OxmlElement("w:fldChar")
    fldEnd.set(qn("w:fldCharType"), "end")
    run._r.append(fldBegin)
    run._r.append(instr)
    run._r.append(fldSep)
    run._r.append(placeholder)
    run._r.append(fldEnd)


def set_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), ACCENT)
    pbdr.append(bottom)
    pPr.append(pbdr)


# ==========================================================================
# Build the document
# ==========================================================================
def build_doc():
    doc = Document()
    set_base_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    # ---------------- Title page ----------------
    for _ in range(3):
        doc.add_paragraph()
    # Banner-style title
    band = doc.add_table(rows=1, cols=1)
    _no_table_borders(band)
    bcell = band.rows[0].cells[0]
    _shade(bcell, PRIMARY)
    _set_cell_margins(bcell, top=400, bottom=400, left=300, right=300)
    tp = bcell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("Finance Research Agent")
    _runfmt(tr, size=30, bold=True, color="FFFFFF")
    sp = bcell.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run("Full End-to-End Workflow \u2014 A Beginner-Friendly Blueprint")
    _runfmt(sr, size=13, color="DCF0EC")

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("From the moment you type a command to the moment a report lands on disk \u2014\nexplained so anyone can follow, with every stage mapped to the real code.")
    _runfmt(r, size=11.5, italic=True, color=SECONDARY)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("Generated June 19, 2026   \u00b7   Word edition of WORKFLOW.md")
    _runfmt(r, size=10.5, color=ACCENT, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    # Disclaimer callout on the title page
    add_callout(
        doc, "Disclaimer (straight from the project)",
        [
            "This software produces educational / research output. It is NOT financial "
            "advice, places NO real orders, and uses FREE public data that may be stale "
            "or wrong. Always verify before acting.",
        ],
        fill=SAND, bar_color=GOLD,
    )

    doc.add_page_break()

    # ---------------- Table of contents ----------------
    add_heading(doc, "Table of Contents", 1)
    add_body(doc, "The headings below are clickable in Word once the field is updated "
                  "(right-click \u2192 Update Field, or press F9).", italic=True,
             size=9.5)
    add_toc_field(doc)
    doc.add_page_break()

    # ---------------- 0. ELI5 ----------------
    add_heading(doc, "0. TL;DR \u2014 Explain It Like I'm 5", 1)
    add_body(doc,
             "Imagine you ask a very careful research team a question like \u201cWhat are "
             "the best IT companies to invest in, in India?\u201d")
    add_callout(
        doc, "The research team, in plain language",
        [
            [("1. A librarian", True), (" figures out which companies you might mean (e.g. the big Indian IT firms) and makes a list.", False)],
            [("2. A calculator robot", True), (" downloads free public numbers (profits, debt, price trends) and scores each company on five qualities \u2014 just math, no opinions \u2014 and prints a \u201chow trustworthy is this data?\u201d sticker.", False)],
            [("3. A panel of AI analysts", True), (" reads those scores and writes short, evidence-based opinions. They are forbidden from making up numbers.", False)],
            [("4. Two AIs debate", True), (" \u2014 one optimistic \u201cBull\u201d, one cautious \u201cBear\u201d.", False)],
            [("5. A risk & tax checker", True), (" adds rules specific to you (e.g. Indian capital-gains tax, or \u201ckeep it low-risk for a German student\u201d).", False)],
            [("6. A manager", True), (" combines everything into a final ranked list and writes a tidy report (Markdown, plus an optional Excel sheet).", False)],
        ],
        fill=MINT, bar_color=ACCENT,
    )
    add_body(doc,
             "You get back a ranked shortlist of stocks with reasons, risks, and tax "
             "notes. Nothing is bought; it's purely research.")

    add_divider(doc)

    # ---------------- 1. Overview ----------------
    add_heading(doc, "1. Overview \u2014 What This Project Does and Why", 1)
    add_body(doc,
             "The Finance Research Agent (FRA) is a command-line program that ranks the "
             "\u201cbest\u201d stocks to research for a given theme (for example, \u201cbest IT stocks "
             "in India\u201d or a single ticker like SAP.DE). It blends deterministic financial "
             "math (reproducible, never guesses) with AI language models (which explain and "
             "contextualize the math). The result is a written investment research report "
             "tailored to one of two investor profiles: an Indian adult investor, or a "
             "German student investor.")
    add_callout(
        doc, "The core idea",
        ["Numbers come from real data and formulas; AI only interprets them. This makes "
         "the output more trustworthy than asking a chatbot to \u201cpick stocks,\u201d because "
         "the AI is never allowed to invent figures."],
        fill=LIGHT, bar_color=PRIMARY,
    )

    # ---------------- 2. Glossary ----------------
    add_heading(doc, "2. Key Concepts / Glossary", 1)
    glossary = [
        ("CLI (Command-Line Interface)", "A program you run by typing a command in a terminal instead of clicking buttons. Here it's python -m src.cli ..."),
        ("Ticker", "A stock's short code, e.g. INFY.NS (Infosys, India NSE) or SAP.DE (SAP, Germany Xetra)."),
        ("Universe", "The full list of candidate companies before narrowing down (e.g. all NIFTY 500 companies)."),
        ("Profile", "A saved set of preferences for one kind of investor (country, currency, tax rules, risk, emphasis). YAML files in config/profiles/."),
        ("Factor", "A measurable \u201cquality\u201d of a company. FRA uses five: Quality, Value, Momentum, Financial Health, Earnings Quality."),
        ("Factor engine", "The deterministic calculator that scores every company on each factor and ranks them. No AI involved."),
        ("Composite score", "A single 0\u20131 number combining the five factors using the profile's weights. Higher = ranks better."),
        ("Percentile rank", "\u201cBetter than X% of the others.\u201d Each company is scored relative to its peers, not on an absolute scale."),
        ("LLM (Large Language Model)", "An AI that reads and writes text (GPT, Claude, or a local Llama via Ollama). Here, LLMs play \u201canalysts.\u201d"),
        ("Agent", "A single specialized worker in the pipeline (e.g. \u201cfundamentals analyst\u201d). Some use an LLM; many are pure math."),
        ("Pipeline / orchestrator", "The conveyor belt that runs the agents in order and passes data between them."),
        ("Snapshot (CompanySnapshot)", "A standardized bundle of one company's data (price, margins, debt, momentum, etc.)."),
        ("Data health card", "A traffic-light summary (OK / WARN / CRITICAL) of how complete and trustworthy this run's data was."),
        ("Debate (Bull vs Bear)", "Two AI researchers argue the optimistic vs cautious case, so the report shows both sides."),
        ("Coverage", "The fraction of expected data fields actually available. Low coverage \u2192 score is \u201cshrunk\u201d toward neutral."),
        ("MCP (Model Context Protocol)", "A standard way for AI tools to call external data/tools. Not currently wired into FRA \u2014 the shipped version uses yfinance/Stooq directly."),
    ]
    styled_table(doc, ["Term", "Plain-language meaning"],
                 [( (t,), m) for t, m in glossary],
                 col_widths=[2.1, 4.3], font_size=9)

    add_divider(doc)

    # ---------------- 3. Architecture ----------------
    add_heading(doc, "3. High-Level Architecture", 1)
    add_body(doc,
             "FRA is a chain of agents. Data starts as your typed question and flows "
             "through each stage, getting richer at every step, until it becomes a report. "
             "A single shared object \u2014 the AgentState (src/graph/state.py) \u2014 is the "
             "\u201cclipboard\u201d that travels through the whole chain, accumulating results.")
    add_image_with_caption(
        doc, ARCH_PNG,
        "Figure 1. Component architecture \u2014 CLI \u2192 orchestrator \u2192 agent pipeline \u2192 outputs, "
        "with external data, LLM provider and profile YAMLs feeding in.")
    add_heading(doc, "Three layers, in plain terms", 3)
    add_bullets(doc, [
        [("Perception (get data): ", True), ("src/data/ \u2014 fetches numbers and news from free sources and caches them on disk.", False)],
        [("Brain (think): ", True), ("src/factors/ (deterministic math) + src/agents/ (AI interpretation + debate).", False)],
        [("Action (write it down): ", True), ("src/report/ (Markdown/Excel) + src/memory/ (a running log of past runs).", False)],
    ])
    add_callout(
        doc, "The whole thing in one line (from the README)",
        ["Universe \u2192 DataProvider \u2192 FactorEngine \u2192 Analysts \u2192 Bull/Bear Debate "
         "\u2192 Risk+Profile \u2192 Manager \u2192 Report"],
        fill=LIGHT, bar_color=SECONDARY,
    )

    doc.add_page_break()

    # ---------------- 4. Workflow ----------------
    add_heading(doc, "4. The End-to-End Workflow, Stage by Stage", 1)
    add_body(doc,
             "The orchestrator runs the stages below. With LangGraph installed it uses a "
             "real graph; otherwise it falls back to a plain sequential runner \u2014 both run "
             "the exact same stages (see _seq_run and _build_langgraph in "
             "src/graph/orchestrator.py).")
    add_image_with_caption(
        doc, FLOW_PNG,
        "Figure 2. Stage-by-stage workflow with the should_run_llm decision branch "
        "(AI path vs deterministic quant-only path).")

    # Stage summary table
    add_heading(doc, "Stages at a glance", 2)
    stage_rows = [
        ("0", "You type a command", "--profile, --target, --top, --no-llm", "cli.py parses options, loads profile, builds AgentState", "AgentState + console banner", "src/cli.py"),
        ("1", "Orchestrator (first stage)", "free-text target, profile", "Detect single ticker vs theme; pick sector & constituents (live, then seeded)", "candidate_tickers, candidate_meta", "src/agents/universe.py"),
        ("2", "After universe built", "candidate tickers + profile", "Fetch snapshots, data-health card, run factor engine \u2192 composite scores", "snapshots, factor_reports, shortlist, data_health, input_hash", "src/agents/quant.py, src/factors/engine.py"),
        ("3", "If LLM enabled & shortlist", "shortlist context", "4 analysts (fundamentals, technical, news, macro) interpret numbers only", "analyst_signals", "src/agents/fundamentals.py, technical.py, news_sentiment.py, macro.py"),
        ("4", "After analysts", "analyst signals + factors", "Bull vs Bear personas argue the 2\u20133 most compelling/concerning names", "debate turns", "src/agents/researchers.py"),
        ("5", "After debate (or quant path)", "picks + profile", "Apply concentration cap, ETF preference, volatility notes, tax notes", "risk_notes, tax_notes", "src/agents/risk_profile.py"),
        ("6", "Final reconciliation", "factors + signals + debate", "Build ranked FinalPicks with thesis, risks, horizon, after-tax estimate", "picks", "src/agents/manager.py"),
        ("7", "After manager", "full AgentState", "Render Markdown (Jinja2) + optional multi-sheet Excel", "report_path, excel_path (files on disk)", "src/report/generator.py"),
        ("8", "After report", "run summary", "Append compact JSON record of the run to history log", ".fra_memory/index.jsonl entry", "src/memory/store.py"),
    ]
    styled_table(
        doc,
        ["#", "Trigger", "Input", "Action", "Output", "Implementing file"],
        stage_rows,
        col_widths=[0.3, 1.15, 1.1, 1.7, 1.2, 1.35], font_size=7.8)

    add_body(doc, "")
    add_callout(
        doc, "Decision point \u2014 should_run_llm",
        ["After Stage 2 the orchestrator calls should_run_llm (src/graph/conditional_logic.py). "
         "If you passed --no-llm, OR the shortlist is empty, it skips all AI stages and jumps to a "
         "deterministic manager (manager.run_quant_only) \u2192 risk \u2192 report. Otherwise it proceeds "
         "through the AI analysts and debate."],
        fill=SAND, bar_color=GOLD,
    )

    # Detailed stage notes
    add_heading(doc, "Stage details", 2)
    add_heading(doc, "Stage 2 \u2014 Quant + Factor Engine (the deterministic calculator)", 3)
    add_body(doc, "This is the analytical heart and the most important stage. For each candidate it:")
    add_bullets(doc, [
        [("Fetches a snapshot ", True), ("via DataProvider.get_snapshot (yfinance primary; cross-checks price against Stooq). All fetches cached in .fra_cache/.", False)],
        [("Flags cross-currency picks ", True), ("(e.g. a US stock for a German investor).", False)],
        [("Applies hard risk constraints ", True), ("from the profile (min market cap, max volatility) via _passes_constraints.", False)],
        [("Builds the Data Health card ", True), ("(build_card in src/data/health.py): tickers fetched, avg coverage, source agreement, OK/WARN/CRITICAL.", False)],
        [("Runs the factor engine ", True), ("(rank_universe): extracts metrics, percentile-ranks each across the universe, averages into factor scores, then combines via profile weights into a composite. Low-coverage composites are shrunk toward 0.5.", False)],
        [("Factor regime check ", True), ("(src/factors/decay.py) flags recently underperforming factors.", False)],
        [("Computes a reproducible input_hash ", True), ("so an identical re-run is verifiable.", False)],
    ])

    add_heading(doc, "Stage 3 \u2014 Analyst Agents (AI interprets the numbers)", 3)
    add_body(doc,
             "Four analysts run in sequence. Each builds a compact data bundle "
             "(shortlist_context in src/agents/_common.py), asks the LLM to reason only over "
             "the provided numbers, and falls back to a deterministic heuristic if the LLM is "
             "unavailable. The shared system prompt (SYSTEM_RULES) forbids inventing numbers.")
    add_bullets(doc, [
        [("3a. Fundamentals: ", True), ("Quality, Value, Financial Health, Earnings Quality; also a best-effort insider-trading signal from SEC EDGAR (US tickers only).", False)],
        [("3b. Technical: ", True), ("interprets price momentum and volatility into a trend stance.", False)],
        [("3c. News + Sentiment: ", True), ("headlines from GDELT (falls back to yfinance news), classifies aggregate sentiment.", False)],
        [("3d. Macro: ", True), ("one short, universe-wide context paragraph for the profile's country/currency.", False)],
    ])

    add_heading(doc, "Stages 4\u20138 \u2014 Debate, Risk, Manager, Report, Memory", 3)
    add_bullets(doc, [
        [("Stage 4 \u2014 Bull vs Bear Debate: ", True), ("two LLM personas argue the optimistic vs cautious case, repeated for --rounds rounds (default 1).", False)],
        [("Stage 5 \u2014 Risk + Profile/Tax: ", True), ("deterministically applies concentration caps, ETF preference, volatility notes, and verbatim tax notes from the profile YAML.", False)],
        [("Stage 6 \u2014 Research Manager: ", True), ("combines everything into ranked FinalPicks with thesis, key risks, confidence, horizon, after-tax estimate; backfills any forgotten shortlisted ticker.", False)],
        [("Stage 7 \u2014 Report Generator: ", True), ("renders Markdown via Jinja2 to reports/<timestamp>-<profile>-<target>.md, plus optional .xlsx.", False)],
        [("Stage 8 \u2014 Memory Log: ", True), ("appends a compact JSON record to .fra_memory/index.jsonl for the history command.", False)],
    ])

    doc.add_page_break()

    # ---------------- 5. Inputs & Outputs ----------------
    add_heading(doc, "5. Inputs & Outputs", 1)
    add_heading(doc, "What you provide", 3)
    styled_table(doc, ["Option", "Meaning"], [
        (("--profile (required)",), "india_adult or germany_student"),
        (("--target (required)",), "a ticker (SAP.DE) or a theme (\u201cbest IT stocks in India\u201d)"),
        (("--top",), "how many picks to return (default 10)"),
        (("--universe / --domain",), "narrow the candidate pool"),
        (("--rounds",), "number of Bull vs Bear debate rounds (default 1)"),
        (("--no-llm",), "deterministic, AI-free run"),
        (("--no-excel",), "skip the Excel workbook"),
        (("--as-of YYYY-MM-DD",), "evaluate as of a past date"),
    ], col_widths=[1.9, 4.5], font_size=9)

    add_heading(doc, "What you get back", 3)
    add_bullets(doc, [
        [("A Markdown report ", True), ("in reports/ and (optionally) an Excel workbook.", False)],
        [("A console summary: ", True), ("a \u201cTop picks\u201d table and a Data Health line.", False)],
        [("A history entry ", True), ("you can later list with python -m src.cli history.", False)],
    ])
    add_callout(
        doc, "What the report contains (real example)",
        ["Header (target, universe, top-N, input hash) + disclaimer; a Data health card "
         "(e.g. OK, 6/6 tickers fetched, 82% coverage); factor regime warnings; Final Picks "
         "(composite score, profile fit, coverage, factor std-dev, confidence, horizon, "
         "after-tax estimate, thesis, key risks, tax notes); a Factor breakdown table; "
         "Analyst signals; the Bull vs Bear transcript; Risk + profile notes; Tax notes; "
         "and a Methodology section listing factor weights."],
        fill=LIGHT, bar_color=PRIMARY,
    )

    # ---------------- 6. Configuration ----------------
    add_heading(doc, "6. Configuration & Setup", 1)
    add_heading(doc, "Prerequisites", 3)
    add_bullets(doc, [
        "Python 3 (modern, with pydantic>=2.7).",
        "Dependencies in requirements.txt (typer, rich, yfinance, pandas, numpy, requests, beautifulsoup4, langgraph, jinja2, openpyxl, plus optional openai/anthropic).",
        "Optional, for local AI: Ollama running a model like llama3.1:8b.",
    ])
    add_heading(doc, "Install & run (Windows / PowerShell)", 3)
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Inches(0.2)
    cr = code.add_run(
        "python -m venv .venv && .venv\\Scripts\\activate\n"
        "pip install -r requirements.txt\n"
        "copy .env.example .env\n\n"
        "# Optional local LLM:\n"
        "ollama pull llama3.1:8b\n"
        "ollama serve\n\n"
        "# Run a research pass:\n"
        "python -m src.cli research --profile india_adult --target \"best IT stocks in India\" --top 10\n\n"
        "# Offline / deterministic (no AI):\n"
        "python -m src.cli research --profile india_adult --target \"best banks in India\" --top 10 --no-llm")
    _runfmt(cr, size=9, color=PRIMARY, font="Consolas")
    _shade_paragraph(code, "F2F7F8")

    add_heading(doc, "Environment variables (.env)", 3)
    styled_table(doc, ["Variable", "Purpose", "Default"], [
        (("LLM_PROVIDER",), "openai, anthropic, ollama, or cursor_io", "ollama"),
        (("LLM_MODEL",), "provider-specific model name", "llama3.1:8b"),
        (("LLM_TEMPERATURE",), "randomness (0 = deterministic)", "0"),
        (("OPENAI_API_KEY / ANTHROPIC_API_KEY",), "keys for hosted LLMs", "empty"),
        (("OLLAMA_HOST",), "local Ollama endpoint", "http://localhost:11434"),
        (("FRA_CACHE_DIR",), "where data is cached", "./.fra_cache"),
    ], col_widths=[2.2, 3.0, 1.2], font_size=9)

    add_body(doc, "")
    add_callout(
        doc, "No API key? It still runs.",
        ["If no LLM is reachable, get_llm() in src/llm/factory.py returns a deterministic "
         "stub, and every analyst falls back to its math-based heuristic. You still get a "
         "valid factor-based report \u2014 just with sparser AI prose."],
        fill=MINT, bar_color=ACCENT,
    )
    add_heading(doc, "Where things live (config-as-data)", 3)
    add_body(doc,
             "Profile YAMLs in config/profiles/ control universe defaults, factor weights, "
             "tax rules, risk constraints, and currency \u2014 so most behavior can be tuned "
             "without editing Python. The two shipped profiles:")
    add_bullets(doc, [
        [("india_adult.yaml ", True), ("\u2014 INR, NIFTY-based, momentum-tilted, Indian capital-gains tax.", False)],
        [("germany_student.yaml ", True), ("\u2014 EUR, DAX/MDAX, quality/health-tilted, lower risk, Abgeltungssteuer.", False)],
    ])

    doc.add_page_break()

    # ---------------- 7. Walkthrough ----------------
    add_heading(doc, "7. Concrete Walkthrough \u2014 One Real Request", 1)
    code2 = doc.add_paragraph()
    code2.paragraph_format.left_indent = Inches(0.2)
    cr2 = code2.add_run(
        'python -m src.cli research --profile india_adult --target "best IT companies in India" --top 10')
    _runfmt(cr2, size=9, color=PRIMARY, font="Consolas")
    _shade_paragraph(code2, "F2F7F8")
    add_bullets(doc, [
        [("1. cli.py ", True), ("loads india_adult.yaml, builds the AgentState, prints the banner, and calls orchestrator.run.", False)],
        [("2. Universe ", True), ("detects the Information Technology sector from \u201cIT\u201d and yields TCS.NS, INFY.NS, TECHM.NS, WIPRO.NS, HCLTECH.NS, LTIM.NS.", False)],
        [("3. Quant ", True), ("fetches a snapshot per ticker (+ Stooq cross-check), builds the Data Health card (OK, 6/6 fetched, 82% coverage, single-source), runs the factor engine, computes input hash 1aceed39273729da.", False)],
        [("4. Analysts ", True), ("read the scores. For TCS: fundamentals bullish (ROE 48.4%, op margin 25.3%, PE 15.3); technical bearish (12-1m momentum \u221231.9%); news bearish ($70m legal hit). Each claim cites a real number.", False)],
        [("5. Debate ", True), ("Bull argues \u201cworld-class franchises after a reset\u201d; Bear argues \u201cthe cycle is still rolling over.\u201d", False)],
        [("6. Risk + Tax ", True), ("adds \u201c10% per-name cap\u201d, \u201cvolatility > 60% filtered\u201d, and Indian LTCG/STCG notes.", False)],
        [("7. Manager ", True), ("produces ranked picks (TCS #1, Infosys #2, \u2026) each with thesis, risks, >12-month horizon, after-tax estimate (e.g. TCS +8.7%). LTIM.NS had 0% coverage \u2192 low-confidence backfilled entry: graceful degradation.", False)],
        [("8. Report + Memory ", True), ("writes reports/20260619-140733-india_adult-best-IT-companies-in-India.md and logs the run. Console prints the Top-picks table, hash, and path.", False)],
    ])

    add_divider(doc)

    # ---------------- 8. Extension points ----------------
    add_heading(doc, "8. Extension Points \u2014 How to Modify Behavior", 1)
    styled_table(doc, ["You want to\u2026", "Change this"], [
        (("Add a new investor profile",), "Add a YAML in config/profiles/ (copy an existing one); auto-discovered by profiles and load_profile."),
        (("Change factor weights / tax / risk limits",), "Edit the relevant config/profiles/*.yaml. No code change needed."),
        (("Add or change a factor",), "Add an extractor in src/factors/metrics.py (register in FACTORS); the engine picks it up."),
        (("Swap the data source (e.g. MCP or paid API)",), "Re-implement the DataProvider interface in src/data/provider.py; the rest is unchanged."),
        (("Add a new analyst",), "Create a module in src/agents/ exposing run(state), then wire it into _seq_run and _build_langgraph."),
        (("Change the report layout",), "Edit src/report/templates/report.md.j2 (Markdown) or src/report/excel.py (Excel)."),
        (("Use a different / no LLM",), "Set LLM_PROVIDER in .env, or pass --no-llm for a deterministic run."),
        (("Tune debate depth",), "Pass --rounds N."),
        (("Add CLI options or commands",), "Edit src/cli.py (it uses typer)."),
    ], col_widths=[2.5, 3.9], font_size=9)

    # ---------------- 9. Reusable template ----------------
    add_heading(doc, "9. Reusable Blueprint Template", 1)
    add_body(doc, "This document followed a repeatable skeleton you can reuse for any "
                  "data-or-agent pipeline:")
    add_bullets(doc, [
        [("TL;DR / ELI5 ", True), ("\u2014 one analogy a non-technical reader can grasp.", False)],
        [("Overview ", True), ("\u2014 what & why in 2\u20133 sentences.", False)],
        [("Glossary ", True), ("\u2014 define every piece of jargon.", False)],
        [("Architecture ", True), ("\u2014 components + a diagram + the 3-layer (perceive/think/act) framing.", False)],
        [("Stage-by-stage workflow ", True), ("\u2014 trigger \u2192 input \u2192 action \u2192 output \u2192 implementing file + a flow diagram.", False)],
        [("Inputs & Outputs ", True), ("\u2014 what the user gives and gets (with a real example).", False)],
        [("Configuration & Setup ", True), ("\u2014 prerequisites, env vars, install/run.", False)],
        [("Concrete walkthrough ", True), ("\u2014 trace one real request end to end.", False)],
        [("Extension points ", True), ("\u2014 a \u201cto change X, edit Y\u201d table.", False)],
    ])

    add_divider(doc)
    fm = doc.add_paragraph()
    r = fm.add_run(
        "File map: src/cli.py (entry) \u00b7 src/graph/orchestrator.py (pipeline) \u00b7 "
        "src/graph/state.py (shared state) \u00b7 src/agents/* (workers) \u00b7 src/factors/* "
        "(deterministic math) \u00b7 src/data/* (data + caching) \u00b7 src/llm/factory.py "
        "(LLM providers) \u00b7 src/report/* (output) \u00b7 src/memory/store.py (history) \u00b7 "
        "config/profiles/*.yaml (settings).")
    _runfmt(r, size=9, italic=True, color=SECONDARY)

    doc.save(DOCX_PATH)
    return DOCX_PATH


def _shade_paragraph(paragraph, hex_fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def main():
    print("Rendering architecture diagram...")
    render_architecture()
    print(f"  -> {ARCH_PNG}")
    print("Rendering workflow flow diagram...")
    render_workflow_flow()
    print(f"  -> {FLOW_PNG}")
    print("Building Word document...")
    path = build_doc()
    size = os.path.getsize(path)
    print(f"  -> {path} ({size/1024:.1f} KB)")
    print("Done.")


if __name__ == "__main__":
    main()
